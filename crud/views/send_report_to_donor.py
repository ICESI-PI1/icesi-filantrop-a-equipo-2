from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required

from crud.models import *
from crud.models import Document as AcademicDocument
from datetime import datetime
from docx import Document
from docx.shared import Pt
import traceback

from crud.views.request_info_update import send_email
from django.contrib.auth import get_user_model



"""
Renders the page to generate and send reports to donors. 

Also, once the user has typed all required information, gets that info and calls the different functions to generate and display the report on page.
"""
@login_required
def send_report_to_donor(request):
    if request.method == 'POST':
        # Dictionary to save the info that will be charged to the html with the generated report.
        html_data = {}

        try:
            # Get the info to generate the report and send the email
            selected_donor_id = request.POST.get('donor-id', '')

            selected_student_id = request.POST.get('student-id', '')

            if selected_donor_id == '' or selected_student_id == '':
                raise ValueError('Error. Seleccione un donante y un estudiante')
                
            email_message = request.POST.get('message-area', '')

            # Charges the info to the dictionary to charge it again with the report
            html_data['selected_donor_id'] = selected_donor_id
            html_data['email_message'] = email_message
            
            date = datetime.now().date()
            semester = f'{date.year} - 1' if date.month < 6 else f'{date.year} - 2'

            report_type = request.POST.get('report-type', '')

            selected_student = Student.objects.get(id=selected_student_id)

            student_testimony = None
            non_academic_activities = None
            crea_assistance = None
            academic_info = None 
            academic_report_name = None

            if 'Reporte general' in report_type:
                student_testimony = request.POST.get('student-testimony', '')

                non_academic_activities = extract_non_academic_information(selected_student.student_code)
                crea_assistance = extract_crea_information(selected_student.student_code)
                academic_info, academic_report_name = extract_academic_information(selected_student_id)

                academic_info = 'No se registra información académica.' if academic_info is None else academic_info
            
                # Charges the academic report name if there is a matching document
                html_data['academic_report_name'] = academic_report_name

                print(f'ACTIVIDADES NO ACADÉMICAS: {non_academic_activities}')
                print(f'CREA: {crea_assistance}')

            elif 'Reporte de actividades no académicas' in report_type:
                non_academic_activities = extract_non_academic_information(selected_student.student_code)

            elif 'Reporte de asistencia al CREA' in report_type:
                crea_assistance = extract_crea_information(selected_student.student_code)

            report_name = generate_report(date, semester, selected_student, report_type, testimony=student_testimony, non_academic_activities=non_academic_activities, crea_assistance=crea_assistance, academic_info=academic_info)
                
            result_message = "Reporte generado con éxito"

            # Charges the generated report name to the html once it has been created successfully
            html_data['generated_report'] = report_name

        # Catches if there is a problem in the generation of the report
        except ValueError as ve:
            result_message = str(ve)

        except Exception as e:
            traceback.print_exc()
            print(f"Error: {e}")

            report_name = None
            result_message = "Error en la generación de reporte"
        
        # Charges the operation's result message to notice the user about it
        html_data['result_message'] = result_message

        return render(request, 'send_report_to_donor.html', html_data)

    if 'Filantropía' in request.user.user_type:
        return render(request, 'send_report_to_donor.html', {
            'user': request.user
        })
    
    else:
        return redirect('/home/')


"""
Searches into the DB for the CREA assistance information of the student passed as parameter, and returns it in a string.
"""
def extract_crea_information(student_code):
    try:
        crea_information = CREAReport.objects.filter(student_code=student_code)

        print(f'INFO: {crea_information}')

        info_str = 'No se registran asistencias a monitorías.' if not crea_information.exists() is None else ''

        for info in crea_information:
            info_str += f'\nNombre monitoría: {info.monitor_name}\n'
            info_str += f'Razón de asistencia: {info.reason}\n'
            info_str += f'Resultado de la monitoría: {info.result}\n'
            info_str += f'Fecha de la monitoría: {info.date}\n'
    
    except Exception as e:
        traceback.print_exc()
        print(f"Error: {e}")

    return info_str


"""
Searches into the DB for the non academic information of the student passed as parameter, and returns it in a string.
"""
def extract_non_academic_information(student_code):
    try:
        non_academic_information = NonAcademicActvitiesReport.objects.filter(student_code=student_code)

        info_str = ''
        info_str = 'No se registran asistencias a actividades no académicas.' if not non_academic_information.exists() is None else ''

        for info in non_academic_information:
            info_str += f'\nNombre actividad: {info.activity}\n'
            info_str += f'Cantidad de horas asistidas: {info.activity_hours}\n'
    
    except Exception as e:
        traceback.print_exc()
        print(f"Error: {e}")

    return info_str


"""
Searches into the DB for the academic information of the student passed as parameter, and returns the name of that report.
"""
def extract_academic_information(student_id):
    try:
        academic_info = AcademicDocument.objects.filter(codigo_estudiante=student_id).first()

        academic_report_name = str(academic_info.uploadedFile).split('/')[-1]

    # Catches the exception when there is not a matching document
    except Exception as e:
        traceback.print_exc()
        print(f"Error: {e}")

        academic_info = None
        academic_report_name = None

    return academic_info, academic_report_name


"""
Manages the event when the user clicks the send button, confirming that she/he is in agreement by sending the generated report.
"""
def send_report(request):
    try:
        # Gets the info
        selected_donor_id = request.POST.get('donor-id2', '')
        selected_donor = Donor.objects.get(id=selected_donor_id)

        email_message = request.POST.get('email-message', '')
        
        # Array of document paths to send in the email
        attachment_paths = []

        generated_report_name = request.POST.get('report-name', '')
        report_path = f'crud/static/reports/{generated_report_name}'
        
        # Appends the path of the report
        attachment_paths.append(report_path)

        academic_report_name = request.POST.get('academic-report-name', '')

        # Checks if the academic_report_name is present and appends its path to the array if it is
        if academic_report_name:
            academic_report_path = f'media/Uploaded Files/{academic_report_name}'
            attachment_paths.append(academic_report_path)

        # This is a generic subject for the email. It can be changed
        subject = 'Reporte general de beneficiario'

        # Sends the email by using the function send_email() on reques_info_update.py
        send_email(selected_donor.email, email_message, subject, attachment_paths=attachment_paths)

        return redirect('/sendReportToDonor/')

    # Notices the user if something went wrong
    except Exception as e:
        traceback.print_exc()
        print(f'Error: {e}')

        result_message = 'Error al enviar email al donante'

        return render(request, 'send_report_to_donor.html', {
            'email_sending_result_message': result_message,
            "generated_report": generated_report_name,
            "selected_donor_id": selected_donor_id,
            "email_message": email_message,
        })


"""
Receives certain information, obtains the basic format of the report (depending on the value of report_type parameter), and fills in the available fields with the information received.

At the final, returns the name of the generated report.
"""
def generate_report(date, semester, student, report_type, testimony=None, non_academic_activities=None, crea_assistance=None, academic_info=None):
    try:
        report_name = f'{report_type} {student.student_code} - {semester}.docx'
        output_path = f'crud/static/reports/{report_name}'

        # Reads the format content
        format_content = read_report_format(report_type)

        # Creates a new doc to add the modifications
        modified_doc = Document()

        for paragraph_info in format_content:
            new_paragraph = modified_doc.add_paragraph()

            # Replaces the fields with the information received as parameters
            paragraph_text = paragraph_info['paragraph_text']
            paragraph_text = paragraph_text.replace("[Fecha]", str(date))
            paragraph_text = paragraph_text.replace("[estudiante]", str(student.name))
            paragraph_text = paragraph_text.replace("[semestre]", semester)

            if 'Reporte general' in report_type:
                paragraph_text = paragraph_text.replace("[testimonio_estudiante]", testimony)
                paragraph_text = paragraph_text.replace("[actividades_no_académicas]", str(non_academic_activities))
                paragraph_text = paragraph_text.replace("[asistencia_monitorias]", str(crea_assistance))

                # If academic_info is None, removes the field available to this purpose. Otherwise, fills that field
                if academic_info is not None:
                    paragraph_text = paragraph_text.replace("[info_academica]", str(academic_info))
                
                else:
                    paragraph_text = paragraph_text.replace("[info_academica]", '')

            elif 'Reporte de actividades no académicas' in report_type:
                paragraph_text = paragraph_text.replace("[actividades_no_académicas]", str(non_academic_activities))

            elif 'Reporte de asistencia al CREA' in report_type:
                paragraph_text = paragraph_text.replace("[asistencia_monitorias]", str(crea_assistance))

            # Adds the modified text to the doc
            new_paragraph.add_run(paragraph_text)

            for run in new_paragraph.runs:
                for run_info in paragraph_info['runs']:
                    run.font.size = Pt(run_info['font_size']) if run_info['font_size'] else None
                    run.font.bold = run_info['font_bold']
                    run.font.italic = run_info['font_italic']
                    run.font.name = run_info['font_name']
                    new_paragraph.alignment = run_info['alignment']

        # Saves the .docx and .pdf of the report
        modified_doc.save(output_path)

        return output_path.split('/')[-1]

    except Exception as e:
        traceback.print_exc()
        print(f'Error por acá: {e}')


"""
Reads a .docx document containing the basic format of the report to be generated. 
"""
def read_report_format(report_type):
    # Checks the type of the report and reads the corresponding base
    if 'Reporte general' in report_type:
        doc = Document('./crud/static/formats/Reporte general beneficiario.docx')

    elif 'Reporte de actividades no académicas' in report_type:
        doc = Document('./crud/static/formats/Reporte de actividades no académicas beneficiario.docx')

    elif 'Reporte de asistencia al CREA' in report_type:
        doc = Document('./crud/static/formats/Reporte de asistencia al CREA beneficiario.docx')

    doc_content = []

    # Builds a new document with the same text of the base format read
    for paragraph in doc.paragraphs:
        para_content = []

        for run in paragraph.runs:
            run_text = run.text
            font_size = run.font.size.pt if run.font.size else None
            font_bold = run.font.bold
            font_italic = run.font.italic
            paragraph_alignment = paragraph.alignment
            font_name = run.font.name

            # Text styles
            para_content.append({
                'text': run_text,
                'font_size': font_size,
                'font_bold': font_bold,
                'font_italic': font_italic,
                'alignment': paragraph_alignment,
                'font_name': font_name,
            })

        doc_content.append({
            'paragraph_text': paragraph.text,
            'runs': para_content,
        })

    return doc_content
