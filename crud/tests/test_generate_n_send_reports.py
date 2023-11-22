from datetime import date
from django.contrib.auth.models import User
from django.test import TestCase
from django.urls import reverse
from unittest.mock import patch, MagicMock
from crud.models import *
from crud.views.send_report_to_donor import *
from datetime import datetime
from docx import Document


class GenerateAndSendReportTestCase(TestCase):

    def setUp(self):
        self.student1 = Student.objects.create(
            student_code='A00456789',
            name='Student1',
            genre='M',
            id_type='CC',
            id_number='1234567890',
            email='ejemplo@gmail.com',
            institutional_email='estudiante@universidad.com',
            icfes_score=300,
            birth_date=datetime.strptime('2000-1-1', '%Y-%m-%d').date(),
            cellphone_number='123456789',
            accumulated_average=4.5,
            credits_studied=120
        )

        self.date1 = datetime.strptime('2023-11-20', '%Y-%m-%d').date()
        self.semester1 = '2023 - 2'
        self.report_type1 = 'Reporte general'
        self.report_type2 = 'Reporte de actividades no académicas'
        self.report_type3 = 'Reporte de asistencia al CREA'
        self.testimony1 = 'Testimonio de ejemplo'
        self.non_academic_activities1 = 'No se registran asistencias a actividades no académicas'
        self.crea_assistance1 = 'No se registran asistencias a monitorías'
        self.academic_info1 = 'No se registra información académica'



    def test_read_general_report(self):
        result = read_report_format(self.report_type1)

        self.assertIsNotNone(result)
        self.assertGreater(len(result), 0)

        title = result[0]["paragraph_text"]
        self.assertIn("Reporte general beneficiario", title)


    def test_read_non_academic_report(self):
        result = read_report_format(self.report_type2)

        self.assertIsNotNone(result)
        self.assertGreater(len(result), 0)

        title = result[0]["paragraph_text"]
        self.assertIn("Reporte de asistencia a actividades no académicas beneficiario", title)


    def test_read_crea_report(self):
        result = read_report_format(self.report_type3)

        self.assertIsNotNone(result)
        self.assertGreater(len(result), 0)

        title = result[0]["paragraph_text"]
        self.assertIn("Reporte de asistencia a monitorías del CREA beneficiario", title)


    def test_generate_general_report(self):
        result = generate_report(self.date1, self.semester1, self.student1, self.report_type1, self.testimony1, self.non_academic_activities1, self.crea_assistance1, self.academic_info1)

        self.assertTrue(result.endswith('.docx'))
        self.assertIn(f'{self.student1.student_code} - {self.semester1}', result)

        report = Document(f'crud/static/reports/{result}')

        title_found = False
        date_found = False
        testimony_found = False
        semester_found = False
        student_found = False

        for paragraph in report.paragraphs:
            if 'Reporte general beneficiario' in paragraph.text:
                title_found = True
            
            if str(self.date1) in paragraph.text:
                date_found = True
            
            if self.testimony1 in paragraph.text:
                testimony_found = True

            if self.semester1 in paragraph.text:
                semester_found = True
            
            if self.student1.name in paragraph.text:
                student_found = True

        self.assertTrue(title_found)
        self.assertTrue(date_found)
        self.assertTrue(testimony_found)
        self.assertTrue(student_found)
        self.assertTrue(semester_found)


    def test_generate_non_academic_report(self):
        result = generate_report(self.date1, self.semester1, self.student1, self.report_type2, self.non_academic_activities1)

        self.assertTrue(result.endswith('.docx'))
        self.assertIn(f'{self.student1.student_code} - {self.semester1}', result)

        report = Document(f'crud/static/reports/{result}')

        title_found = False
        date_found = False
        semester_found = False
        student_found = False

        for paragraph in report.paragraphs:
            if 'Reporte de asistencia a actividades no académicas beneficiario' in paragraph.text:
                title_found = True
            
            if str(self.date1) in paragraph.text:
                date_found = True

            if self.semester1 in paragraph.text:
                semester_found = True
            
            if self.student1.name in paragraph.text:
                student_found = True

        self.assertTrue(title_found)
        self.assertTrue(date_found)
        self.assertTrue(student_found)
        self.assertTrue(semester_found)


    def test_generate_crea_report(self):
        result = generate_report(self.date1, self.semester1, self.student1, self.report_type3, self.crea_assistance1)

        self.assertTrue(result.endswith('.docx'))
        self.assertIn(f'{self.student1.student_code} - {self.semester1}', result)

        report = Document(f'crud/static/reports/{result}')

        title_found = False
        date_found = False
        semester_found = False
        student_found = False

        for paragraph in report.paragraphs:
            if 'Reporte de asistencia a monitorías del CREA beneficiario' in paragraph.text:
                title_found = True
            
            if str(self.date1) in paragraph.text:
                date_found = True

            if self.semester1 in paragraph.text:
                semester_found = True
            
            if self.student1.name in paragraph.text:
                student_found = True

        self.assertTrue(title_found)
        self.assertTrue(date_found)
        self.assertTrue(student_found)
        self.assertTrue(semester_found)